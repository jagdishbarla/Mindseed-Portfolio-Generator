import os
import pandas as pd
import requests
from io import BytesIO
from docx import Document

# Load all necessary sheets from Google Drive Excel
def load_data(google_drive_url):
    response = requests.get(google_drive_url)
    if response.status_code != 200:
        raise Exception("Failed to download file from Google Drive")

    file_data = BytesIO(response.content)

    # Load May & June learning levels
    may_df = pd.read_excel(file_data, sheet_name="May")
    june_df = pd.read_excel(file_data, sheet_name="June")

    # Reload to reset pointer, then load all subject description sheets
    file_data.seek(0)
    desc_df = pd.read_excel(file_data, sheet_name=None)  # Load all sheets into a dictionary

    return may_df, june_df, desc_df


# Map grade to relevant subjects
def get_subjects_by_grade(grade):
    grade = str(grade).strip()
    if grade == "Toddler":
        return ["Pre Math", "Fine Motor"]
    elif grade == "Pre-K":
        return ["Pre Math", "Numeracy", "Phonemic Awareness", "Fine Motor"]
    elif grade == "K1":
        return ["Pattern Writing", "Numeracy", "Phonemic Awareness"]
    elif grade == "K2":
        return ["Writing", "Numeracy", "Phonemic Awareness", "Reading", "Pattern Writing"]
    return []


# Generate child-voice narrative for a subject
def generate_subject_narrative(subject, level_may, level_june, desc_df):
    try:
        subject_desc = desc_df[subject]
        row = subject_desc[subject_desc["Level"] == level_june].iloc[0]
        learning = row["What I Was Learning"]
        why = row["Why It Matters"]
        apply = row["Real-World Application"]
    except:
        return f"Could not generate learning summary for {subject}."

    if level_june > level_may:
        return (f"I moved from Level {level_may} to {level_june} in {subject}, learning to {learning.lower()}. "
                f"It matters because {why.lower()}, and I use it when I {apply.lower()}.")
    elif level_june == level_may:
        return (f"I’m continuing at Level {level_may} in {subject}, working on {learning.lower()}. "
                f"It matters because {why.lower()}, and I’m getting better every day.")
    else:
        return (f"I'm revisiting Level {level_june} in {subject} to strengthen how I {learning.lower()}. "
                f"It helps because {why.lower()}, and I use it when I {apply.lower()}.")


# Create the actual DOCX portfolio for one child
def create_portfolio(child_row_may, child_row_june, desc_df, output_path):
    doc = Document()
    name = child_row_june["Child Name"]
    school = child_row_june["School"]
    grade = child_row_june["Grade"]

    # Header
    doc.add_heading(f"My Learning Portfolio – {name}", 0)
    doc.add_paragraph(f"School: {school}")
    doc.add_paragraph(f"Grade: {grade}")

    subjects = get_subjects_by_grade(grade)

    for subject in subjects:
        try:
            level_may = int(child_row_may[subject])
            level_june = int(child_row_june[subject])
            doc.add_heading(f"My {subject} Journey", level=1)
            narrative = generate_subject_narrative(subject, level_may, level_june, desc_df)
            doc.add_paragraph(narrative)
        except Exception as e:
            doc.add_paragraph(f"{subject}: No data available.")

    filename = f"{name.replace(' ', '_')}_Portfolio.docx"
    filepath = os.path.join(output_path, filename)
    doc.save(filepath)
    return filepath
